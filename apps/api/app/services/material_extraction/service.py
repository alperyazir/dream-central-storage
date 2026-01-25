"""Service for extracting text from teacher materials (PDF, TXT, DOCX)."""

from __future__ import annotations

import logging
from io import BytesIO
from typing import TYPE_CHECKING, Callable

import fitz  # PyMuPDF for PDF extraction

from app.core.config import get_settings
from app.services.minio import get_minio_client
from app.services.material_extraction.models import (
    ExtractionMethod,
    FileType,
    MaterialExtractionError,
    MaterialExtractionResult,
    PageText,
    UnsupportedFileTypeError,
    ExtractionFailedError,
)

if TYPE_CHECKING:
    from minio import Minio

logger = logging.getLogger(__name__)


class MaterialExtractionService:
    """Service for extracting text from teacher materials."""

    def __init__(self, minio_client: Minio | None = None) -> None:
        """Initialize the service.

        Args:
            minio_client: Optional MinIO client (created if not provided)
        """
        self._client = minio_client

    @property
    def client(self) -> Minio:
        """Get MinIO client, creating if needed."""
        if self._client is None:
            settings = get_settings()
            self._client = get_minio_client(settings)
        return self._client

    async def extract_material_text(
        self,
        material_id: int,
        teacher_id: str,
        material_name: str,
        file_type: str,
        progress_callback: Callable[[int, int], None] | None = None,
    ) -> MaterialExtractionResult:
        """Extract text from a teacher material.

        Args:
            material_id: Database ID of the material
            teacher_id: Teacher ID (folder name in storage)
            material_name: Material filename
            file_type: File extension (pdf, txt, docx)
            progress_callback: Optional callback for progress updates

        Returns:
            MaterialExtractionResult with extracted text

        Raises:
            UnsupportedFileTypeError: If file type is not supported
            ExtractionFailedError: If extraction fails
        """
        file_type_enum = FileType.from_extension(file_type)

        if file_type_enum is None or not file_type_enum.is_text_extractable:
            raise UnsupportedFileTypeError(
                f"File type '{file_type}' is not supported for text extraction",
                material_name,
            )

        logger.info(
            "Starting text extraction for material %s (teacher: %s, type: %s)",
            material_name,
            teacher_id,
            file_type_enum.value,
        )

        # Download the material file
        file_data = self._download_material(teacher_id, material_name)

        # Extract based on file type
        if file_type_enum == FileType.PDF:
            return await self._extract_pdf(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_data=file_data,
                progress_callback=progress_callback,
            )
        elif file_type_enum == FileType.TXT:
            return await self._extract_txt(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_data=file_data,
                progress_callback=progress_callback,
            )
        elif file_type_enum == FileType.DOCX:
            return await self._extract_docx(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_data=file_data,
                progress_callback=progress_callback,
            )
        else:
            raise UnsupportedFileTypeError(
                f"File type '{file_type}' extraction not implemented",
                material_name,
            )

    def _download_material(self, teacher_id: str, material_name: str) -> bytes:
        """Download material file from MinIO.

        Args:
            teacher_id: Teacher ID
            material_name: Material filename

        Returns:
            File contents as bytes

        Raises:
            ExtractionFailedError: If file cannot be downloaded
        """
        settings = get_settings()
        bucket = settings.minio_teachers_bucket
        object_key = f"{teacher_id}/materials/{material_name}"

        try:
            response = self.client.get_object(bucket, object_key)
            data = response.read()
            response.close()
            response.release_conn()
            return data
        except Exception as e:
            logger.error("Failed to download material '%s': %s", object_key, e)
            raise ExtractionFailedError(
                f"Failed to download file: {e}",
                material_name,
            ) from e

    async def _extract_pdf(
        self,
        material_id: int,
        teacher_id: str,
        material_name: str,
        file_data: bytes,
        progress_callback: Callable[[int, int], None] | None = None,
    ) -> MaterialExtractionResult:
        """Extract text from PDF file.

        Args:
            material_id: Material database ID
            teacher_id: Teacher ID
            material_name: Material filename
            file_data: PDF file bytes
            progress_callback: Optional progress callback

        Returns:
            MaterialExtractionResult with extracted pages
        """
        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
        except Exception as e:
            logger.error("Failed to open PDF '%s': %s", material_name, e)
            raise ExtractionFailedError(
                f"Failed to open PDF: {e}",
                material_name,
            ) from e

        try:
            total_pages = len(doc)
            pages: list[PageText] = []
            total_word_count = 0

            for page_num in range(total_pages):
                page = doc[page_num]
                text = self._extract_pdf_page_text(page)

                page_text = PageText(
                    page_number=page_num + 1,
                    text=text,
                    method=ExtractionMethod.NATIVE,
                )
                pages.append(page_text)
                total_word_count += page_text.word_count

                if progress_callback:
                    progress_callback(page_num + 1, total_pages)

            logger.info(
                "PDF extraction complete: %d pages, %d words",
                total_pages,
                total_word_count,
            )

            return MaterialExtractionResult(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_type=FileType.PDF,
                total_pages=total_pages,
                total_word_count=total_word_count,
                method=ExtractionMethod.NATIVE,
                pages=pages,
            )
        finally:
            doc.close()

    def _extract_pdf_page_text(self, page: fitz.Page) -> str:
        """Extract text from a PDF page with layout handling.

        Args:
            page: PyMuPDF page object

        Returns:
            Extracted text
        """
        try:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            text_blocks = [b for b in blocks if b.get("type") == 0]

            if not text_blocks:
                return ""

            # Sort by vertical then horizontal position
            sorted_blocks = sorted(
                text_blocks,
                key=lambda b: (round(b["bbox"][1] / 20) * 20, b["bbox"][0]),
            )

            text_parts: list[str] = []
            for block in sorted_blocks:
                block_lines: list[str] = []
                for line in block.get("lines", []):
                    spans_text = []
                    for span in line.get("spans", []):
                        span_text = span.get("text", "")
                        if span_text:
                            spans_text.append(span_text)
                    line_text = "".join(spans_text).strip()
                    if line_text:
                        block_lines.append(line_text)
                if block_lines:
                    text_parts.append("\n".join(block_lines))

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning("Failed to extract text blocks: %s, falling back", e)
            return page.get_text("text")

    async def _extract_txt(
        self,
        material_id: int,
        teacher_id: str,
        material_name: str,
        file_data: bytes,
        progress_callback: Callable[[int, int], None] | None = None,
    ) -> MaterialExtractionResult:
        """Extract text from TXT file.

        Args:
            material_id: Material database ID
            teacher_id: Teacher ID
            material_name: Material filename
            file_data: TXT file bytes
            progress_callback: Optional progress callback

        Returns:
            MaterialExtractionResult with extracted text
        """
        try:
            # Try UTF-8 first, then fallback to latin-1
            try:
                text = file_data.decode("utf-8")
            except UnicodeDecodeError:
                text = file_data.decode("latin-1")

            word_count = len(text.split())

            if progress_callback:
                progress_callback(1, 1)

            logger.info(
                "TXT extraction complete: %d words",
                word_count,
            )

            return MaterialExtractionResult(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_type=FileType.TXT,
                total_pages=1,
                total_word_count=word_count,
                method=ExtractionMethod.DIRECT,
                pages=[
                    PageText(
                        page_number=1,
                        text=text,
                        method=ExtractionMethod.DIRECT,
                        word_count=word_count,
                    )
                ],
            )
        except Exception as e:
            logger.error("Failed to extract TXT '%s': %s", material_name, e)
            raise ExtractionFailedError(
                f"Failed to read text file: {e}",
                material_name,
            ) from e

    async def _extract_docx(
        self,
        material_id: int,
        teacher_id: str,
        material_name: str,
        file_data: bytes,
        progress_callback: Callable[[int, int], None] | None = None,
    ) -> MaterialExtractionResult:
        """Extract text from DOCX file.

        Args:
            material_id: Material database ID
            teacher_id: Teacher ID
            material_name: Material filename
            file_data: DOCX file bytes
            progress_callback: Optional progress callback

        Returns:
            MaterialExtractionResult with extracted text
        """
        try:
            from docx import Document
        except ImportError as e:
            logger.error("python-docx not installed")
            raise ExtractionFailedError(
                "DOCX extraction requires python-docx package",
                material_name,
            ) from e

        try:
            doc = Document(BytesIO(file_data))

            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            text = "\n\n".join(paragraphs)
            word_count = len(text.split())

            if progress_callback:
                progress_callback(1, 1)

            logger.info(
                "DOCX extraction complete: %d paragraphs, %d words",
                len(paragraphs),
                word_count,
            )

            return MaterialExtractionResult(
                material_id=material_id,
                teacher_id=teacher_id,
                material_name=material_name,
                file_type=FileType.DOCX,
                total_pages=1,
                total_word_count=word_count,
                method=ExtractionMethod.NATIVE,
                pages=[
                    PageText(
                        page_number=1,
                        text=text,
                        method=ExtractionMethod.NATIVE,
                        word_count=word_count,
                    )
                ],
            )
        except Exception as e:
            logger.error("Failed to extract DOCX '%s': %s", material_name, e)
            raise ExtractionFailedError(
                f"Failed to read DOCX file: {e}",
                material_name,
            ) from e


# Module-level service instance getter
_service: MaterialExtractionService | None = None


def get_material_extraction_service() -> MaterialExtractionService:
    """Get or create the material extraction service singleton."""
    global _service
    if _service is None:
        _service = MaterialExtractionService()
    return _service
